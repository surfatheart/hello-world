# 创新医药政策质性分析与主题分析（DeepSeek API）
# 用法1：在 Jupyter 新建一个代码单元格，把本文件全部内容粘进去，运行即可。
# 用法2：命令行执行  python 创新医药政策分析.py
# 运行前：①下方 pip 安装依赖（首次）②在“全局配置”处填密钥与文件夹路径。

# ===== 首次运行请先安装依赖（去掉行首#执行一次）=====
# pip install openai pymupdf python-docx pandas matplotlib openpyxl tqdm

import os
import re
import json
import time
import glob
import hashlib
from pathlib import Path

import fitz                       # PyMuPDF，解析 PDF
from docx import Document         # 解析 DOCX
import pandas as pd
import numpy as np
import matplotlib
import matplotlib.pyplot as plt
from matplotlib import font_manager
from tqdm import tqdm
from openai import OpenAI         # DeepSeek 兼容 OpenAI SDK

print("依赖库导入完成。")

# ====================================================================
# ## 2. 全局配置（**只需要改这一格**）
#
# 把下面三处改成你自己的即可：DeepSeek 密钥、政策文件夹路径、输出目录。
# ====================================================================

# ====== ① DeepSeek 密钥（把 sk-xxx 换成你的）======
DEEPSEEK_API_KEY = "sk-在此填入你的DeepSeek密钥"
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
DEEPSEEK_MODEL = "deepseek-chat"          # 也可改成 "deepseek-reasoner"

# ====== ② 政策文件夹路径（你的 D 盘师姐政策文件夹）======
# Windows 路径前面加 r，例如 r"D:\师姐政策文件夹"
POLICY_DIR = r"D:\师姐政策文件夹"

# ====== ③ 结果输出文件夹（会自动创建）======
OUTPUT_DIR = r"D:\师姐政策文件夹\政策分析输出"

# ---- 下面不用改 ----
os.makedirs(OUTPUT_DIR, exist_ok=True)
CACHE_PATH = os.path.join(OUTPUT_DIR, "coding_cache_policy.json")  # 编码缓存文件
print("政策文件夹:", POLICY_DIR)
print("输出目录  :", OUTPUT_DIR)

# ====== 1.3 评估维度框架（概念驱动，预设维度）======
# 这些是“先验维度”，来自创新医药政策评价的常见框架。
# 同时允许“数据驱动”的新增维度：模型在编码时可提出文本充分支持的新维度。
DIMENSIONS = [
    "创新要素配置",   # 平台、人才、资源、种质资源库、算力等
    "创新方向",       # 新靶点/新机制/前沿技术；中药新药/经典名方/院内制剂等
    "临床前研究",     # 化合物筛选/AI药物设计；道地药材质量/炮制工艺/物质基准等
    "临床研究",       # 国际多中心/随机双盲；三结合证据体系/真实世界研究等
    "成果转化",       # 专利转让/作价入股/许可交易；院内制剂→中药新药等
    "资金支持",       # 研发投入补助、分阶段资助、专项资金等
    "政策支持",       # 审评审批加速、入院绿色通道、医保支付改革、调剂使用等
]

# 每个维度的释义（用于提示词，帮助模型准确归类）
DIMENSION_DEFINITIONS = {
    "创新要素配置": "创新所需的平台、人才、资源等要素布局。化学药/生物制品如生物样本库、CRO/CDMO平台、AI算力平台；中医药如道地药材种质资源库、中药炮制技术传承基地、区域制剂中心、中药全产业链大模型。",
    "创新方向": "鼓励的创新类型与重点领域。化学药/生物制品如新靶点、新机制、细胞/基因治疗、mRNA等前沿技术；中医药如中药新药、经典名方开发、名优中成药二次开发、院内制剂转化。",
    "临床前研究": "上市前的基础与工艺研究。化学药/生物制品如新靶点发现、化合物筛选、AI药物设计；中医药如道地药材质量研究、炮制工艺规范、经典名方物质基准研究、中药标准化。",
    "临床研究": "临床试验与证据生成。化学药/生物制品如国际多中心、随机双盲对照试验；中医药如‘三结合’证据体系（中医药理论、人用经验、临床试验）、真实世界研究。",
    "成果转化": "科研成果向产品/产业的转化路径。如专利转让、作价入股、许可交易；中医药以院内制剂→中药新药为主要路径，辅以名医验方开发、中药大品种二次开发。",
    "资金支持": "财政/基金对研发与转化的资金投入，如分阶段研发投入补助、最高额度、专项资助等。",
    "政策支持": "制度与流程层面的支持，如审评审批加速、入院绿色通道、医保支付改革、中药特色审评、院内制剂调剂使用放宽、配方颗粒标准制定等。",
}

# ====== 1.4 药物类型轴 ======
DRUG_TYPES = ["化学药与生物制品", "中医药", "通用"]  # 通用=两类共用/未明确区分

# ====== 1.5 省份清单与元数据（用于识别归属 + 报告参考画像）======
# value 中的关键词用于从“PDF标题（文件名+文档开头大标题）”优先识别省份归属，其次看正文。
# 已覆盖全国 31 个省级行政区，关键词含省名、简称、省会及主要城市/园区，尽量减少“未识别”。
PROVINCE_KEYWORDS = {
    # —— 重点 10 省（含产业园区关键词）——
    "江苏省": ["江苏", "苏州", "南京", "泰州", "无锡", "常州", "连云港", "苏州工业园"],
    "上海市": ["上海", "张江", "浦东", "临港"],
    "广东省": ["广东", "广州", "深圳", "珠海", "佛山", "东莞", "坪山", "生物岛", "粤港澳"],
    "北京市": ["北京", "中关村", "亦庄", "昌平", "大兴"],
    "浙江省": ["浙江", "杭州", "余杭", "宁波", "温州", "嘉兴", "医药港"],
    "湖北省": ["湖北", "武汉", "光谷", "宜昌", "襄阳"],
    "山东省": ["山东", "济南", "青岛", "烟台", "潍坊", "齐鲁"],
    "四川省": ["四川", "成都", "天府", "绵阳"],
    "天津市": ["天津", "滨海新区", "滨海"],
    "安徽省": ["安徽", "合肥", "芜湖", "蚌埠", "科大硅谷"],
    # —— 其余省 / 自治区 / 直辖市 ——
    "重庆市": ["重庆", "两江新区"],
    "河北省": ["河北", "石家庄", "雄安", "唐山"],
    "山西省": ["山西", "太原"],
    "辽宁省": ["辽宁", "沈阳", "大连"],
    "吉林省": ["吉林", "长春"],
    "黑龙江省": ["黑龙江", "哈尔滨"],
    "福建省": ["福建", "福州", "厦门", "泉州"],
    "江西省": ["江西", "南昌", "赣州"],
    "河南省": ["河南", "郑州", "洛阳"],
    "湖南省": ["湖南", "长沙", "株洲"],
    "海南省": ["海南", "海口", "三亚", "博鳌"],
    "贵州省": ["贵州", "贵阳"],
    "云南省": ["云南", "昆明"],
    "陕西省": ["陕西", "西安", "杨凌"],
    "甘肃省": ["甘肃", "兰州"],
    "青海省": ["青海", "西宁"],
    "内蒙古自治区": ["内蒙古", "呼和浩特", "包头"],
    "广西壮族自治区": ["广西", "南宁", "桂林"],
    "宁夏回族自治区": ["宁夏", "银川"],
    "新疆维吾尔自治区": ["新疆", "乌鲁木齐"],
    "西藏自治区": ["西藏", "拉萨"],
}

# 国家级文件的识别关键词（当标题中无省份时兜底，避免误标“未识别”）
NATIONAL_KEYWORDS = [
    "国务院", "国办发", "国发", "国家医保局", "国家药监局", "国家卫生健康委",
    "国家中医药管理局", "国家发展改革委", "工业和信息化部", "全国", "国家层面",
]

# 省份产业画像（仅作报告背景参考，不替代政策文本证据；报告中引用须注明为背景资料）
PROVINCE_PROFILE = {
    "江苏省": "规上营收全国第一，创新药获批连续两年居首，苏州/南京/泰州园区稳居全国前十。",
    "上海市": "全链条生态最完善，CDE所在地，张江药谷集聚全球20大药企中19家，2025年产业规模近万亿。",
    "广东省": "医疗器械全国领先，基因测序/呼吸机超40%份额，广州生物岛+深圳坪山双核驱动。",
    "北京市": "科研与制度资源最密集，中科院/协和/301+CDE支撑原始创新，First-in-class项目最多。",
    "浙江省": "‘生物医药+数字医疗’融合突出，杭州医药港+余杭未来科技城，融资与临床批件活跃。",
    "湖北省": "武汉光谷生物城全国第六，2024年生命健康规模超5500亿，中部龙头。",
    "山东省": "化学药与原料药基础强，齐鲁制药等龙头支撑，2023年营收第二。",
    "四川省": "成都高新区全国园区第三，西部创新转化高地，科伦/康诺亚等崛起。",
    "天津市": "制造与协同优势突出，滨海新区为亚洲最大胰岛素基地，康希诺等领军。",
    "安徽省": "合肥‘科大硅谷’+芜湖/蚌埠联动，AI制药与基因治疗加速，近年创新增速亮眼。",
}

# 参考对照矩阵（用户提供的“化学药与生物制品 vs 中医药”先验框架）。
# 仅作为对照表的“参考列”，真实结论以政策文本编码为准。
REFERENCE_MATRIX = {
    "创新要素配置": {"化学药与生物制品": "生物样本库、CRO/CDMO平台、AI算力平台",
                 "中医药": "道地药材种质资源库、中药炮制技术传承基地、区域制剂中心、中药全产业链大模型"},
    "创新方向":   {"化学药与生物制品": "新靶点、新机制、前沿技术（细胞/基因治疗、mRNA）",
                 "中医药": "中药新药、经典名方开发、名优中成药二次开发、院内制剂转化"},
    "临床前研究": {"化学药与生物制品": "新靶点发现、化合物筛选、AI药物设计",
                 "中医药": "道地药材质量研究、炮制工艺规范、经典名方物质基准研究、中药标准化"},
    "临床研究":   {"化学药与生物制品": "国际多中心临床试验、随机双盲对照试验",
                 "中医药": "‘三结合’证据体系（中医药理论、人用经验、临床试验）、真实世界研究"},
    "成果转化":   {"化学药与生物制品": "专利转让、作价入股、许可交易",
                 "中医药": "院内制剂→中药新药为主，辅以名医验方开发、中药大品种二次开发"},
    "资金支持":   {"化学药与生物制品": "1类新药分阶段给予研发投入40%、最高3000万元支持",
                 "中医药": "通常与新药同等支持，部分地方对院内制剂转化、经典名方等有专项资助"},
    "政策支持":   {"化学药与生物制品": "审评审批加速、入院绿色通道、医保支付改革",
                 "中医药": "中药特色审评体系、院内制剂调剂使用放宽、中药配方颗粒标准制定"},
}

# ====================================================================
# ## 3. 文本抽取
# 支持 PDF、DOCX、DOC（尽力）、TXT。自动识别省份归属。
# ====================================================================

def extract_pdf(path):
    text = []
    with fitz.open(path) as doc:
        for page in doc:
            text.append(page.get_text())
    return "\n".join(text)

def extract_docx(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    # 表格文本也抽取
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)

def extract_txt(path):
    for enc in ("utf-8", "gbk", "gb18030", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    return ""

def extract_text(path):
    ext = Path(path).suffix.lower()
    try:
        if ext == ".pdf":
            return extract_pdf(path)
        if ext == ".docx":
            return extract_docx(path)
        if ext in (".txt", ".md"):
            return extract_txt(path)
        if ext == ".doc":
            # 旧版 .doc 需要 antiword/libreoffice；这里尝试 textract，失败则跳过
            try:
                import textract
                return textract.process(path).decode("utf-8", "ignore")
            except Exception:
                print(f"[跳过] 无法解析旧版 .doc（建议先转为 .docx）：{os.path.basename(path)}")
                return ""
    except Exception as e:
        print(f"[错误] 解析失败 {os.path.basename(path)}: {e}")
        return ""
    return ""

def get_doc_title(text, max_lines=5):
    """提取“文档标题区”：正文开头的前若干非空行（政策文件的大标题通常在最前面，
    有时第一行是文号，故合并前几行一起作为标题区，识别更稳）。"""
    lines = [l.strip() for l in (text or "").splitlines() if l.strip()]
    return " ".join(lines[:max_lines])[:200]

def detect_province(filename, text):
    """识别省份归属：优先用“PDF标题”（文件名 + 文档开头大标题），其次正文前部，
    再次按国家级文件兜底，仍无则返回‘未识别’。"""
    title_zone = f"{filename or ''} {get_doc_title(text)}"   # 标题区（最高优先级）
    body_zone = (text or "")[:3000]                          # 正文前部（次优先级）

    def best_match(zone):
        hits = {}
        for prov, kws in PROVINCE_KEYWORDS.items():
            score = sum(zone.count(kw) for kw in kws)
            if score:
                hits[prov] = score
        return max(hits, key=hits.get) if hits else None

    # 1) 先从标题区识别（文件名/文档大标题里通常直接写明“XX省/市”）
    prov = best_match(title_zone)
    if prov:
        return prov
    # 2) 标题没有，再看正文前部
    prov = best_match(body_zone)
    if prov:
        return prov
    # 3) 国家级文件兜底
    if any(k in title_zone or k in body_zone for k in NATIONAL_KEYWORDS):
        return "国家级"
    return "未识别"

# ====== 2.1 收集并抽取所有政策文本 ======
PATTERNS = ["*.pdf", "*.docx", "*.doc", "*.txt", "*.md"]
files = []
for pat in PATTERNS:
    files.extend(glob.glob(os.path.join(POLICY_DIR, "**", pat), recursive=True))
files = sorted(set(files))
print(f"在 {POLICY_DIR} 共发现 {len(files)} 个候选文件")

documents = {}   # filename -> {"path","text","province","n_chars"}
for fp in files:
    name = os.path.basename(fp)
    txt = extract_text(fp)
    if len(txt.strip()) < 80:   # 过滤空文件/解析失败
        continue
    documents[name] = {
        "path": fp,
        "text": txt,
        "province": detect_province(name, txt),
        "n_chars": len(txt),
    }
print(f"成功抽取有效文本的文件数：{len(documents)}")
for n, d in documents.items():
    print(f"  - {n}  | 省份={d['province']} | 字数={d['n_chars']}")

# ====================================================================
# ## 4. DeepSeek 编码（开放编码 + 选择性编码）
#
# 对每份文件，按维度抽取编码要点，**每条要点必须附政策原文逐字引证**，并标注药物类型。
# 长文本自动分块后合并。结果写入缓存，避免重复调用。
# ====================================================================

client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)

def load_cache():
    if os.path.exists(CACHE_PATH):
        with open(CACHE_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_cache(cache):
    with open(CACHE_PATH, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)

def chunk_text(text, size=9000, overlap=400):
    """按字符分块，保留少量重叠以免切断语义。"""
    if len(text) <= size:
        return [text]
    chunks, i = [], 0
    while i < len(text):
        chunks.append(text[i:i + size])
        i += size - overlap
    return chunks

def call_deepseek_json(system_prompt, user_prompt, max_retries=4, temperature=0.2):
    """调用 DeepSeek 并强制 JSON 输出，带指数退避重试。"""
    for attempt in range(max_retries):
        try:
            resp = client.chat.completions.create(
                model=DEEPSEEK_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=temperature,
                response_format={"type": "json_object"},
            )
            return json.loads(resp.choices[0].message.content)
        except Exception as e:
            wait = 2 ** attempt
            print(f"  [重试 {attempt+1}/{max_retries}] {e}，{wait}s 后重试…")
            time.sleep(wait)
    raise RuntimeError("DeepSeek 调用多次失败")

CODING_SYSTEM = (
    "你是一名严谨的卫生政策质性研究编码员，擅长政策文本分析与扎根理论编码。"
    "你只依据给定文本进行编码，绝不臆测或补充文本中不存在的信息。"
    "每一条编码都必须能在原文中找到支撑句，并以逐字引证（quote）形式给出。"
)

def build_coding_prompt(doc_name, province, chunk_text_):
    dim_desc = "\n".join(f"- {d}：{DIMENSION_DEFINITIONS[d]}" for d in DIMENSIONS)
    return f"""请对下面这份政策文本片段进行质性编码。文件名：《{doc_name}》（疑似省份：{province}）。

【预设评估维度及释义】
{dim_desc}

【任务】
1. 针对每个预设维度，抽取该片段中**明确涉及**的政策要点；每条要点必须给出原文逐字引证。
2. 为每条要点标注药物类型，取值仅限：{ " / ".join(DRUG_TYPES) }（无法区分时填“通用”）。
3. 若某维度在该片段中**未涉及**，请给出空数组 []，不要编造。
4. 如发现预设维度无法覆盖、但文本充分支持的新主题，放入 "新增维度"。

【输出 JSON 格式（严格遵守）】
{{
  "维度": {{
    "创新要素配置": [{{"编码": "简明要点", "原文": "逐字引证（可截取关键句）", "药物类型": "通用"}}],
    "创新方向": [],
    "临床前研究": [],
    "临床研究": [],
    "成果转化": [],
    "资金支持": [],
    "政策支持": []
  }},
  "新增维度": [{{"维度名": "...", "编码": "...", "原文": "...", "药物类型": "..."}}]
}}

【待编码文本片段】
\"\"\"{chunk_text_}\"\"\"
"""

def merge_codings(codings):
    """合并同一文件多个分块的编码结果。"""
    merged = {"维度": {d: [] for d in DIMENSIONS}, "新增维度": []}
    seen = set()
    for c in codings:
        for d in DIMENSIONS:
            for item in c.get("维度", {}).get(d, []) or []:
                key = (d, item.get("编码", "").strip())
                if key in seen or not item.get("编码"):
                    continue
                seen.add(key)
                merged["维度"][d].append(item)
        for item in c.get("新增维度", []) or []:
            key = ("新增", item.get("维度名", ""), item.get("编码", "").strip())
            if key in seen or not item.get("编码"):
                continue
            seen.add(key)
            merged["新增维度"].append(item)
    return merged

def code_document(doc_name, doc, cache, force=False):
    if (doc_name in cache) and not force:
        return cache[doc_name]
    chunks = chunk_text(doc["text"])
    results = []
    for idx, ch in enumerate(chunks):
        print(f"    编码分块 {idx+1}/{len(chunks)} …")
        prompt = build_coding_prompt(doc_name, doc["province"], ch)
        results.append(call_deepseek_json(CODING_SYSTEM, prompt))
    merged = merge_codings(results)
    merged["_province"] = doc["province"]
    cache[doc_name] = merged
    save_cache(cache)
    return merged

# ====== 4.1 执行编码（带缓存）======
cache = load_cache()
for name, doc in tqdm(documents.items(), desc="编码文件"):
    if name in cache:
        continue
    print(f"\n>>> 正在编码：{name}")
    try:
        code_document(name, doc, cache)
    except Exception as e:
        print(f"  [失败] {name}: {e}")
print("\n编码完成。已缓存到:", CACHE_PATH)

# ====================================================================
# ## 5. 质性编码汇总表
# 生成长表（文件 × 维度 × 编码 × 原文 × 药物类型 × 省份），并导出 Excel。
# ====================================================================

rows = []
for name, doc in documents.items():
    coded = cache.get(name)
    if not coded:
        continue
    prov = doc["province"]
    for d in DIMENSIONS:
        for item in coded.get("维度", {}).get(d, []) or []:
            rows.append({
                "文件": name, "省份": prov, "维度": d,
                "编码要点": item.get("编码", ""),
                "原文引证": item.get("原文", ""),
                "药物类型": item.get("药物类型", "通用"),
            })
    for item in coded.get("新增维度", []) or []:
        rows.append({
            "文件": name, "省份": prov,
            "维度": f"[新增] {item.get('维度名','')}",
            "编码要点": item.get("编码", ""),
            "原文引证": item.get("原文", ""),
            "药物类型": item.get("药物类型", "通用"),
        })

coding_df = pd.DataFrame(rows)
print("编码总条数:", len(coding_df))
coding_df.head(20)

coding_xlsx = os.path.join(OUTPUT_DIR, "附录A_质性编码汇总表.xlsx")
coding_df.to_excel(coding_xlsx, index=False)
print("已导出:", coding_xlsx)

# ====================================================================
# ## 6. 主题分析（跨文件聚类提炼主题）
#
# 将每个维度下、所有文件的编码要点汇总后交给 DeepSeek 进行**轴心编码/主题提炼**，
# 归纳为若干核心主题，并标注覆盖省份、支持文件数与代表性引文。
# ====================================================================

THEME_SYSTEM = (
    "你是质性研究主题分析专家，熟悉 Braun & Clarke 的反身性主题分析。"
    "你只能基于给定的编码条目进行归纳聚类，不得引入外部信息或编造省份/政策。"
)

def build_theme_prompt(dimension, items):
    payload = [{"文件": it["文件"], "省份": it["省份"], "编码": it["编码要点"]} for it in items]
    return f"""以下是“{dimension}”维度下、来自多份省级创新医药政策文本的全部编码要点（JSON 数组）。
请进行主题分析：把语义相近的编码聚成若干（建议 3-6 个）**核心主题**。

要求：
1. 每个主题给出：主题名称、主题阐释（1-3句）、覆盖省份列表、支持文件数、2-4条代表性编码原句（从输入中挑选，不得改写为不存在的内容）。
2. 主题必须忠实于输入编码，禁止臆造。
3. 严格输出 JSON：
{{"主题": [{{"主题名称": "...", "主题阐释": "...", "覆盖省份": ["..."], "支持文件数": 0, "代表编码": ["..."]}}]}}

【输入编码】
{json.dumps(payload, ensure_ascii=False)}
"""

theme_results = {}
for d in DIMENSIONS:
    items = [r for r in rows if r["维度"] == d]
    if len(items) < 2:
        theme_results[d] = {"主题": []}
        continue
    print(f">>> 主题分析：{d}（{len(items)} 条编码）")
    theme_results[d] = call_deepseek_json(THEME_SYSTEM, build_theme_prompt(d, items))

# 保存主题结果
with open(os.path.join(OUTPUT_DIR, "主题分析结果.json"), "w", encoding="utf-8") as f:
    json.dump(theme_results, f, ensure_ascii=False, indent=2)
print("主题分析完成。")

# ====== 6.1 主题分析表 ======
theme_rows = []
for d in DIMENSIONS:
    for t in theme_results.get(d, {}).get("主题", []) or []:
        theme_rows.append({
            "维度": d,
            "主题名称": t.get("主题名称", ""),
            "主题阐释": t.get("主题阐释", ""),
            "覆盖省份": "、".join(t.get("覆盖省份", []) or []),
            "支持文件数": t.get("支持文件数", 0),
            "代表编码": " || ".join(t.get("代表编码", []) or []),
        })
theme_df = pd.DataFrame(theme_rows)
theme_xlsx = os.path.join(OUTPUT_DIR, "主题分析表.xlsx")
theme_df.to_excel(theme_xlsx, index=False)
print("已导出:", theme_xlsx)
theme_df

# ====================================================================
# ## 7. 比较矩阵
# （1）省份 × 维度 编码数量矩阵；（2）化学药/生物制品 vs 中医药 对照表（含参考列）。
# ====================================================================

# ====== 7.1 省份 × 维度 编码数量矩阵 ======
prov_dim = (coding_df[coding_df["维度"].isin(DIMENSIONS)]
            .pivot_table(index="省份", columns="维度", values="编码要点",
                         aggfunc="count", fill_value=0))
# 维度列排序
prov_dim = prov_dim.reindex(columns=DIMENSIONS, fill_value=0)
matrix_xlsx = os.path.join(OUTPUT_DIR, "省份_维度_编码矩阵.xlsx")
prov_dim.to_excel(matrix_xlsx)
print("已导出:", matrix_xlsx)
prov_dim

# ====== 7.2 化学药/生物制品 vs 中医药 对照表 ======
def summarize_codes(dim, drug_type, max_items=6):
    sub = coding_df[(coding_df["维度"] == dim) & (coding_df["药物类型"] == drug_type)]
    pts = list(dict.fromkeys(sub["编码要点"].tolist()))  # 去重保序
    return "；".join(pts[:max_items]) if pts else "（现有政策文本未涉及）"

cmp_rows = []
for d in DIMENSIONS:
    cmp_rows.append({
        "评估维度": d,
        "参考-化学药与生物制品": REFERENCE_MATRIX.get(d, {}).get("化学药与生物制品", ""),
        "参考-中医药": REFERENCE_MATRIX.get(d, {}).get("中医药", ""),
        "实证-化学药与生物制品": summarize_codes(d, "化学药与生物制品"),
        "实证-中医药": summarize_codes(d, "中医药"),
        "实证-通用": summarize_codes(d, "通用"),
    })
compare_df = pd.DataFrame(cmp_rows)
compare_xlsx = os.path.join(OUTPUT_DIR, "化学药生物制品_vs_中医药_对照表.xlsx")
compare_df.to_excel(compare_xlsx, index=False)
print("已导出:", compare_xlsx)
compare_df

# ====================================================================
# ## 8. 可视化
# 中文字体：脚本会尝试常见中文字体；若图中中文显示为方框，请安装并指定 SimHei/Microsoft YaHei。
# ====================================================================

# 中文字体自适应
def set_chinese_font():
    candidates = ["Microsoft YaHei", "SimHei", "PingFang SC", "Heiti SC",
                  "WenQuanYi Zen Hei", "Noto Sans CJK SC", "Source Han Sans SC", "Arial Unicode MS"]
    available = {f.name for f in font_manager.fontManager.ttflist}
    for c in candidates:
        if c in available:
            matplotlib.rcParams["font.sans-serif"] = [c]
            matplotlib.rcParams["axes.unicode_minus"] = False
            print("使用中文字体:", c)
            return
    print("[警告] 未找到中文字体，图中中文可能显示为方框。请安装 SimHei 等字体。")
    matplotlib.rcParams["axes.unicode_minus"] = False
set_chinese_font()

FIG_DIR = os.path.join(OUTPUT_DIR, "图")
os.makedirs(FIG_DIR, exist_ok=True)

# ====== 图1：省份 × 维度 编码数量热力图 ======
if not prov_dim.empty:
    fig, ax = plt.subplots(figsize=(11, max(4, 0.6 * len(prov_dim) + 2)))
    data = prov_dim.values
    im = ax.imshow(data, cmap="YlOrRd", aspect="auto")
    ax.set_xticks(range(len(prov_dim.columns)))
    ax.set_xticklabels(prov_dim.columns, rotation=30, ha="right")
    ax.set_yticks(range(len(prov_dim.index)))
    ax.set_yticklabels(prov_dim.index)
    for i in range(data.shape[0]):
        for j in range(data.shape[1]):
            ax.text(j, i, int(data[i, j]), ha="center", va="center", fontsize=9)
    ax.set_title("各省份在不同评估维度上的政策编码数量分布")
    fig.colorbar(im, ax=ax, label="编码条数")
    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, "图1_省份维度热力图.png"), dpi=200)
    plt.show()

# ====== 图2：各维度编码总数条形图 ======
dim_counts = (coding_df[coding_df["维度"].isin(DIMENSIONS)]
              .groupby("维度").size().reindex(DIMENSIONS, fill_value=0))
fig, ax = plt.subplots(figsize=(10, 5))
ax.bar(dim_counts.index, dim_counts.values, color="#4C72B0")
ax.set_title("各评估维度的政策编码总数")
ax.set_ylabel("编码条数")
plt.xticks(rotation=25, ha="right")
for i, v in enumerate(dim_counts.values):
    ax.text(i, v, str(int(v)), ha="center", va="bottom")
plt.tight_layout()
plt.savefig(os.path.join(FIG_DIR, "图2_维度编码总数.png"), dpi=200)
plt.show()

# ====== 图3：化学药/生物制品 vs 中医药 各维度编码对比 ======
pivot_drug = (coding_df[coding_df["维度"].isin(DIMENSIONS)]
              .pivot_table(index="维度", columns="药物类型", values="编码要点",
                           aggfunc="count", fill_value=0).reindex(DIMENSIONS, fill_value=0))
for col in DRUG_TYPES:
    if col not in pivot_drug.columns:
        pivot_drug[col] = 0
pivot_drug = pivot_drug[DRUG_TYPES]
fig, ax = plt.subplots(figsize=(11, 5))
x = np.arange(len(pivot_drug.index)); w = 0.27
colors = {"化学药与生物制品": "#C44E52", "中医药": "#55A868", "通用": "#8172B3"}
for k, col in enumerate(DRUG_TYPES):
    ax.bar(x + (k - 1) * w, pivot_drug[col].values, w, label=col, color=colors[col])
ax.set_xticks(x); ax.set_xticklabels(pivot_drug.index, rotation=25, ha="right")
ax.set_ylabel("编码条数"); ax.set_title("化学药与生物制品 vs 中医药：各维度政策着力点对比")
ax.legend()
plt.tight_layout()
plt.savefig(os.path.join(FIG_DIR, "图3_化药生物制品_vs_中医药.png"), dpi=200)
plt.show()

# ====== 图4：主题频次条形图（按支持文件数）======
if not theme_df.empty:
    td = theme_df.copy()
    td["标签"] = td["维度"] + " | " + td["主题名称"]
    td = td.sort_values("支持文件数", ascending=True).tail(15)
    fig, ax = plt.subplots(figsize=(11, max(4, 0.45 * len(td) + 2)))
    ax.barh(td["标签"], td["支持文件数"], color="#DD8452")
    ax.set_xlabel("支持文件数")
    ax.set_title("核心主题的支持强度（Top 主题）")
    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, "图4_主题频次.png"), dpi=200)
    plt.show()

print("图表已保存到:", FIG_DIR)

# ====================================================================
# ## 9. 生成约 8000 字研究报告
#
# 仅向模型提供**已抽取的真实编码、主题与矩阵**作为证据，分章节生成以保证篇幅与质量。
# 强制要求：凡证据不足处标注“现有政策文本未涉及”，引用观点须标注来源文件/省份，禁止杜撰。
# ====================================================================

# ====== 9.1 构建“证据包”（提供给模型的唯一事实来源）======
def build_evidence_pack():
    lines = []
    lines.append("【一、纳入文本清单】")
    for name, doc in documents.items():
        lines.append(f"- 《{name}》 省份={doc['province']} 字数={doc['n_chars']}")

    lines.append("\n【二、省份×维度 编码数量矩阵】")
    lines.append(prov_dim.to_string())

    lines.append("\n【三、各维度编码要点（含省份与药物类型；这是唯一可引用的事实证据）】")
    for d in DIMENSIONS:
        lines.append(f"\n## 维度：{d}")
        sub = [r for r in rows if r["维度"] == d]
        for r in sub:
            q = (r["原文引证"] or "")[:120]
            lines.append(f"  - [{r['省份']}|{r['药物类型']}] {r['编码要点']}（原文：{q}）")

    lines.append("\n【四、主题分析结果】")
    for d in DIMENSIONS:
        for t in theme_results.get(d, {}).get("主题", []) or []:
            lines.append(f"  - [{d}] 主题：{t.get('主题名称','')}；省份：{'、'.join(t.get('覆盖省份',[]))}；"
                         f"支持文件数：{t.get('支持文件数',0)}；阐释：{t.get('主题阐释','')}")

    lines.append("\n【五、化学药/生物制品 vs 中医药 实证对照（来自编码）】")
    for _, r in compare_df.iterrows():
        lines.append(f"  - {r['评估维度']}：化学药/生物制品=[{r['实证-化学药与生物制品']}]；"
                     f"中医药=[{r['实证-中医药']}]；通用=[{r['实证-通用']}]")

    lines.append("\n【六、省份产业画像（背景参考资料，非政策文本证据，引用须注明为背景）】")
    for p, prof in PROVINCE_PROFILE.items():
        lines.append(f"  - {p}：{prof}")
    return "\n".join(lines)

EVIDENCE = build_evidence_pack()
print("证据包字符数:", len(EVIDENCE))
print(EVIDENCE[:1500])

REPORT_SYSTEM = (
    "你是一名资深卫生政策研究学者，撰写规范的中文学术研究报告。"
    "你必须严格基于用户提供的【证据包】写作：所有事实性陈述都要能在证据包中找到依据，"
    "引用具体政策要点时标注来源省份/文件；证据不足之处必须写明‘现有政策文本未涉及’，"
    "严禁编造政策条款、数字、省份或文件。语言客观、学术、有逻辑层次。"
)

# 8000 字分章节，给出各章目标字数
REPORT_SECTIONS = [
    ("摘要与关键词", "撰写中文摘要（目的/方法/主要发现/意义，约400字）和5个关键词。", 450),
    ("一、引言", "问题提出、研究背景（可用省份产业画像作背景）、研究意义与研究问题。", 1000),
    ("二、研究方法", "说明质性分析（开放编码+选择性编码）与反身性主题分析相结合的设计、DeepSeek辅助编码流程、文本来源与遴选、编码框架（7个预设维度+数据驱动新增维度）、防杜撰与效度保障（逐字引证、缓存可复现）。", 1100),
    ("三、各评估维度的横向比较", "依次分析创新要素配置、创新方向、临床前研究、临床研究、成果转化、资金支持、政策支持7个维度；每个维度对比省份差异，并对照化学药/生物制品与中医药两条路径。务必引用证据包中的具体编码与省份。", 2400),
    ("四、主题分析发现", "呈现主题分析提炼出的核心主题及其跨省分布与政策意涵。", 1100),
    ("五、省域格局与化药/中医药路径差异讨论", "讨论区域格局、化学药生物制品与中医药政策逻辑差异及其成因。", 1000),
    ("六、政策建议与研究局限", "基于发现提出可操作的政策建议；说明研究局限（样本范围、文本时效、编码主观性等）与展望。", 950),
]

def generate_section(title, instruction, target_words, prev_titles):
    user = f"""【证据包】
{EVIDENCE}

【写作任务】
现在撰写研究报告的章节：「{title}」。
内容要求：{instruction}
目标字数：约 {target_words} 字（中文）。
已写章节（避免重复）：{prev_titles}

请直接输出该章节的正文（含本级标题），使用 Markdown。务必基于证据包，标注来源省份/文件，证据不足处标注‘现有政策文本未涉及’。"""
    resp = client.chat.completions.create(
        model=DEEPSEEK_MODEL,
        messages=[{"role": "system", "content": REPORT_SYSTEM},
                  {"role": "user", "content": user}],
        temperature=0.5,
    )
    return resp.choices[0].message.content

report_parts = ["# 我国省域创新医药政策的质性与主题分析研究\n"]
prev = []
for title, instruction, tgt in REPORT_SECTIONS:
    print(f">>> 生成章节：{title}（目标约{tgt}字）")
    sec = generate_section(title, instruction, tgt, "；".join(prev))
    report_parts.append(sec)
    prev.append(title)
    time.sleep(1)

report_md = "\n\n".join(report_parts)
report_path = os.path.join(OUTPUT_DIR, "创新医药政策研究报告.md")
with open(report_path, "w", encoding="utf-8") as f:
    f.write(report_md)

approx_words = len(re.findall(r"[\u4e00-\u9fff]", report_md))
print(f"\n报告已生成：{report_path}")
print(f"中文字数（汉字计）约：{approx_words}")

# ====================================================================
# ## 10. 汇总产出清单
# 运行结束后，所有结果都在 `政策分析输出/` 目录下。
# ====================================================================

print("=" * 60)
print("产出清单（目录：%s）" % OUTPUT_DIR)
print("=" * 60)
for f in sorted(glob.glob(os.path.join(OUTPUT_DIR, "**", "*"), recursive=True)):
    if os.path.isfile(f):
        print("  -", os.path.relpath(f, OUTPUT_DIR))
print("\n说明：")
print("  · 附录A_质性编码汇总表.xlsx —— 逐条编码 + 原文引证（防杜撰证据）")
print("  · 主题分析表.xlsx / 主题分析结果.json —— 主题分析产出")
print("  · 省份_维度_编码矩阵.xlsx —— 比较矩阵")
print("  · 化学药生物制品_vs_中医药_对照表.xlsx —— 双路径对照")
print("  · 图/*.png —— 可视化图表")
print("  · 创新医药政策研究报告.md —— 约8000字研究报告")
print("  · coding_cache_policy.json —— 编码缓存（重复运行直接复用）")
