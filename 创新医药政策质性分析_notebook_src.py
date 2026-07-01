# %% [markdown]
# # 创新医药政策文本的质性分析与主题分析（DeepSeek API）
#
# 本笔记本对一个文件夹下的所有政策 PDF/DOCX 文本，采用 **整合式混合方法**——
# **框架式质性分析（演绎：沿7个先验维度+政策工具类型做结构化编码与横向比较）** 与
# **归纳式主题分析（Braun & Clarke 反身性主题分析：跨维度提炼潜在主题）** 相结合，
# 并通过 **联合展示（joint display）做三角互证**——调用 **DeepSeek API** 完成，最终产出：
#
# 1. **三级质性编码汇总表**（一级初始概念→二级主轴范畴→三级核心范畴/政策工具 + 业务维度 + 药物类型 + 原文引证 + 理论化判定依据）
# 2. **跨维度主题分析表 + 整合联合展示表**（主题 × 维度，体现两方法互补而非重复）
# 3. **省份 × 维度比较矩阵**、**化学药/生物制品 vs 中医药 对照表**、**政策工具类型分布（表2）**
# 4. **可视化图表 图1~图7**（编码热力图、维度条形图、双路径对比、主题强度、工具分布、整合联合展示等）
# 5. **研究报告**：Markdown 版（图表已内嵌）+ **Word .docx 版（图表与附表内嵌，适合投稿核心期刊）**
#
# **两种方法如何共处（整合研究）**：质性分析是“结构层骨架”（描述各维度有什么、多少、在哪），
# 主题分析是“解释层叙事”（跨维度的潜在模式与趋势）；二者通过联合展示矩阵对应整合、相互印证，
# 互补而非重复。
#
# **防杜撰设计**：每一条编码都要求 DeepSeek 给出政策原文的逐字引证（quote），若文本未涉及
# 某维度则明确标注“未提及”；报告生成阶段仅向模型提供已抽取的真实编码与引文，并强制要求
# 凡证据不足之处标注“现有政策文本未涉及”，不得编造数据、政策条款或省份。
#
# > 运行顺序：从上到下依次执行各单元格即可。第一次运行会调用 API 并写入本地缓存
# > `coding_cache_policy.json`，重复运行将直接读取缓存，避免重复计费。

# %% [markdown]
# ## 0. 安装依赖
# 若已安装可跳过本单元格。

# %%
# !pip install -q openai pymupdf python-docx pandas matplotlib openpyxl tqdm

# %% [markdown]
# ## 1. 导入所有依赖库
# 按惯例，全部 import 集中放在最上面这一格。

# %%
import os
import re
import json
import time
import glob
import hashlib
from pathlib import Path

import fitz                       # PyMuPDF，解析 PDF
from docx import Document         # 解析 DOCX
import pandas as pd
import numpy as np
import matplotlib
import matplotlib.pyplot as plt
from matplotlib import font_manager
from tqdm import tqdm
from openai import OpenAI         # DeepSeek 兼容 OpenAI SDK

print("依赖库导入完成。")

# %% [markdown]
# ## 2. 全局配置（**只需要改这一格**）
#
# 把下面三处改成你自己的即可：DeepSeek 密钥、政策文件夹路径、输出目录。

# %%
# ====== ① DeepSeek 密钥（把 sk-xxx 换成你的）======
DEEPSEEK_API_KEY = "sk-在此填入你的DeepSeek密钥"
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
DEEPSEEK_MODEL = "deepseek-chat"          # 也可改成 "deepseek-reasoner"

# ====== ② 政策文件夹路径（你的 D 盘师姐政策文件夹）======
# Windows 路径前面加 r，例如 r"D:\师姐政策文件夹"
POLICY_DIR = r"D:\师姐政策文件夹"

# ====== ③ 结果输出文件夹（会自动创建）======
OUTPUT_DIR = r"D:\师姐政策文件夹\政策分析输出"

# ---- 下面不用改 ----
os.makedirs(OUTPUT_DIR, exist_ok=True)
CACHE_PATH = os.path.join(OUTPUT_DIR, "coding_cache_policy_v2.json")  # 编码缓存文件（三级编码新版）
print("政策文件夹:", POLICY_DIR)
print("输出目录  :", OUTPUT_DIR)

# %%
# ====== 1.3 评估维度框架（概念驱动，预设维度）======
# 这些是“先验维度”，来自创新医药政策评价的常见框架。
# 同时允许“数据驱动”的新增维度：模型在编码时可提出文本充分支持的新维度。
DIMENSIONS = [
    "创新要素配置",   # 平台、人才、资源、种质资源库、算力等
    "创新方向",       # 新靶点/新机制/前沿技术；中药新药/经典名方/院内制剂等
    "临床前研究",     # 化合物筛选/AI药物设计；道地药材质量/炮制工艺/物质基准等
    "临床研究",       # 国际多中心/随机双盲；三结合证据体系/真实世界研究等
    "成果转化",       # 专利转让/作价入股/许可交易；院内制剂→中药新药等
    "生产与市场准入", # 生产制造、上市许可、集采挂网、入院与推广应用
]
# 业务维度（X轴：产业链/创新链环节）。与“政策工具维度”正交，二者共同构成双维度矩阵。
BUSINESS_DIMENSIONS = DIMENSIONS   # 语义别名，强调这是“业务/产业链”维度

# 每个业务维度的释义（用于提示词，帮助模型准确归类）
DIMENSION_DEFINITIONS = {
    "创新要素配置": "创新所需的平台、人才、资源等要素布局。化学药/生物制品如生物样本库、CRO/CDMO平台、AI算力平台；中医药如道地药材种质资源库、中药炮制技术传承基地、区域制剂中心、中药全产业链大模型。",
    "创新方向": "鼓励的创新类型与重点领域。化学药/生物制品如新靶点、新机制、细胞/基因治疗、mRNA等前沿技术；中医药如中药新药、经典名方开发、名优中成药二次开发、院内制剂转化。",
    "临床前研究": "上市前的基础与工艺研究。化学药/生物制品如新靶点发现、化合物筛选、AI药物设计；中医药如道地药材质量研究、炮制工艺规范、经典名方物质基准研究、中药标准化。",
    "临床研究": "临床试验与证据生成。化学药/生物制品如国际多中心、随机双盲对照试验；中医药如‘三结合’证据体系（中医药理论、人用经验、临床试验）、真实世界研究。",
    "成果转化": "科研成果向产品/产业的转化路径。如专利转让、作价入股、许可交易；中医药以院内制剂→中药新药为主要路径，辅以名医验方开发、中药大品种二次开发。",
    "生产与市场准入": "产业化与市场环节：生产制造、质量体系、上市许可、集采挂网、医保准入、入院使用与市场推广等。",
}

# ====== 1.4 药物类型轴 ======
DRUG_TYPES = ["化学药与生物制品", "中医药", "通用"]  # 通用=两类共用/未明确区分

# ====== 1.4b 政策工具维度（Y轴：三级“核心范畴”，采用 Rothwell & Zegveld 分类）======
# 说明：业务维度=“政策作用于产业链的哪个环节”；政策工具维度=“政府用什么手段作用”。
# 二者划分标准不同、相互正交，共同构成“业务×工具”双维度分析矩阵（见第8节）。
POLICY_TOOLS = {
    "供给型工具": {"定义": "政府从供给侧直接扩大创新要素供给（人才、资金、技术、平台、信息、基础设施），对创新形成推动力。",
                "关键词": ["平台", "人才", "培训", "实验室", "基础设施", "公共服务", "信息共享", "研发投入", "资金投入", "种质资源库", "传承基地", "中心"]},
    "环境型工具": {"定义": "政府通过营造政策环境（目标规划、财税金融、法规管制、知识产权、标准与策略性措施）间接激励创新。",
                "关键词": ["规划", "方案", "补贴", "奖励", "奖补", "税收", "贷款", "金融", "基金", "法规", "条例", "管理办法", "知识产权", "标准", "审评审批", "医保", "价格", "专项资金"]},
    "需求型工具": {"定义": "政府通过拉动市场需求（政府采购、示范应用、贸易管制、市场准入、服务外包）为创新提供市场空间。",
                "关键词": ["采购", "示范", "应用推广", "入院", "挂网", "绿色通道", "市场准入", "调剂使用", "首购", "首用", "订单", "海外"]},
}
TOOL_TYPES = list(POLICY_TOOLS.keys())   # 三级核心范畴：供给型 / 环境型 / 需求型
# 工具细类（二级“主轴范畴”的工具侧参考，McDonnell & Elmore 五类）
TOOL_SUBTYPES = ["命令性工具", "激励性工具", "能力建设工具", "权威重组工具", "劝告性工具"]

# 关键词兜底分类器（当 LLM 未给出合法三级工具时使用）
def classify_tool_by_keyword(text):
    scores = {t: sum(text.count(k) for k in info["关键词"]) for t, info in POLICY_TOOLS.items()}
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "环境型工具"   # 政策文件多为环境型，作默认

# ====== 1.5 省份清单与元数据（用于识别归属 + 报告参考画像）======
# value 中的关键词用于从“PDF标题（文件名+文档开头大标题）”优先识别省份归属，其次看正文。
# 已覆盖全国 31 个省级行政区，关键词含省名、简称、省会及主要城市/园区，尽量减少“未识别”。
PROVINCE_KEYWORDS = {
    # —— 重点 10 省（含产业园区关键词）——
    "江苏省": ["江苏", "苏州", "南京", "泰州", "无锡", "常州", "连云港", "苏州工业园"],
    "上海市": ["上海", "张江", "浦东", "临港"],
    "广东省": ["广东", "广州", "深圳", "珠海", "佛山", "东莞", "坪山", "生物岛", "粤港澳"],
    "北京市": ["北京", "中关村", "亦庄", "昌平", "大兴"],
    "浙江省": ["浙江", "杭州", "余杭", "宁波", "温州", "嘉兴", "医药港"],
    "湖北省": ["湖北", "武汉", "光谷", "宜昌", "襄阳"],
    "山东省": ["山东", "济南", "青岛", "烟台", "潍坊", "齐鲁"],
    "四川省": ["四川", "成都", "天府", "绵阳"],
    "天津市": ["天津", "滨海新区", "滨海"],
    "安徽省": ["安徽", "合肥", "芜湖", "蚌埠", "科大硅谷"],
    # —— 其余省 / 自治区 / 直辖市 ——
    "重庆市": ["重庆", "两江新区"],
    "河北省": ["河北", "石家庄", "雄安", "唐山"],
    "山西省": ["山西", "太原"],
    "辽宁省": ["辽宁", "沈阳", "大连"],
    "吉林省": ["吉林", "长春"],
    "黑龙江省": ["黑龙江", "哈尔滨"],
    "福建省": ["福建", "福州", "厦门", "泉州"],
    "江西省": ["江西", "南昌", "赣州"],
    "河南省": ["河南", "郑州", "洛阳"],
    "湖南省": ["湖南", "长沙", "株洲"],
    "海南省": ["海南", "海口", "三亚", "博鳌"],
    "贵州省": ["贵州", "贵阳"],
    "云南省": ["云南", "昆明"],
    "陕西省": ["陕西", "西安", "杨凌"],
    "甘肃省": ["甘肃", "兰州"],
    "青海省": ["青海", "西宁"],
    "内蒙古自治区": ["内蒙古", "呼和浩特", "包头"],
    "广西壮族自治区": ["广西", "南宁", "桂林"],
    "宁夏回族自治区": ["宁夏", "银川"],
    "新疆维吾尔自治区": ["新疆", "乌鲁木齐"],
    "西藏自治区": ["西藏", "拉萨"],
}

# 国家级文件的识别关键词（当标题中无省份时兜底，避免误标“未识别”）
NATIONAL_KEYWORDS = [
    "国务院", "国办发", "国发", "国家医保局", "国家药监局", "国家卫生健康委",
    "国家中医药管理局", "国家发展改革委", "工业和信息化部", "全国", "国家层面",
]

# 省份产业画像（仅作报告背景参考，不替代政策文本证据；报告中引用须注明为背景资料）
PROVINCE_PROFILE = {
    "江苏省": "规上营收全国第一，创新药获批连续两年居首，苏州/南京/泰州园区稳居全国前十。",
    "上海市": "全链条生态最完善，CDE所在地，张江药谷集聚全球20大药企中19家，2025年产业规模近万亿。",
    "广东省": "医疗器械全国领先，基因测序/呼吸机超40%份额，广州生物岛+深圳坪山双核驱动。",
    "北京市": "科研与制度资源最密集，中科院/协和/301+CDE支撑原始创新，First-in-class项目最多。",
    "浙江省": "‘生物医药+数字医疗’融合突出，杭州医药港+余杭未来科技城，融资与临床批件活跃。",
    "湖北省": "武汉光谷生物城全国第六，2024年生命健康规模超5500亿，中部龙头。",
    "山东省": "化学药与原料药基础强，齐鲁制药等龙头支撑，2023年营收第二。",
    "四川省": "成都高新区全国园区第三，西部创新转化高地，科伦/康诺亚等崛起。",
    "天津市": "制造与协同优势突出，滨海新区为亚洲最大胰岛素基地，康希诺等领军。",
    "安徽省": "合肥‘科大硅谷’+芜湖/蚌埠联动，AI制药与基因治疗加速，近年创新增速亮眼。",
}

# 参考对照（用户提供的先验框架；仅作对照表“参考列”，真实结论以政策文本编码为准）
# —— 业务维度层面的 化药/中医药 对照 ——
REFERENCE_BUSINESS = {
    "创新要素配置": {"化学药与生物制品": "生物样本库、CRO/CDMO平台、AI算力平台",
                 "中医药": "道地药材种质资源库、中药炮制技术传承基地、区域制剂中心、中药全产业链大模型"},
    "创新方向":   {"化学药与生物制品": "新靶点、新机制、前沿技术（细胞/基因治疗、mRNA）",
                 "中医药": "中药新药、经典名方开发、名优中成药二次开发、院内制剂转化"},
    "临床前研究": {"化学药与生物制品": "新靶点发现、化合物筛选、AI药物设计",
                 "中医药": "道地药材质量研究、炮制工艺规范、经典名方物质基准研究、中药标准化"},
    "临床研究":   {"化学药与生物制品": "国际多中心临床试验、随机双盲对照试验",
                 "中医药": "‘三结合’证据体系（中医药理论、人用经验、临床试验）、真实世界研究"},
    "成果转化":   {"化学药与生物制品": "专利转让、作价入股、许可交易",
                 "中医药": "院内制剂→中药新药为主，辅以名医验方开发、中药大品种二次开发"},
    "生产与市场准入": {"化学药与生物制品": "集采挂网、医保准入、入院使用、政府采购",
                   "中医药": "院内制剂调剂使用、中药饮片/配方颗粒流通与推广"},
}
# —— 政策工具维度（Rothwell & Zegveld）层面的 化药/中医药 对照（由用户提供的资金/政策条目按工具理论重组）——
REFERENCE_TOOL = {
    "供给型工具": {"化学药与生物制品": "生物样本库、CRO/CDMO平台、AI算力平台、研发人才与资金投入",
                "中医药": "道地药材种质资源库、炮制技术传承基地、区域制剂中心、全产业链大模型"},
    "环境型工具": {"化学药与生物制品": "1类新药分阶段研发投入补助（如40%、最高3000万元）、审评审批加速、医保支付改革",
                "中医药": "院内制剂/经典名方专项资助、中药特色审评体系、中药配方颗粒标准制定"},
    "需求型工具": {"化学药与生物制品": "入院绿色通道、集采挂网、政府采购",
                "中医药": "院内制剂调剂使用放宽、中药应用推广与示范"},
}
REFERENCE_MATRIX = REFERENCE_BUSINESS   # 兼容别名（业务维度对照）

# %% [markdown]
# ## 3. 文本抽取
# 支持 PDF、DOCX、DOC（尽力）、TXT。自动识别省份归属。

# %%
def extract_pdf(path):
    text = []
    with fitz.open(path) as doc:
        for page in doc:
            text.append(page.get_text())
    return "\n".join(text)

def extract_docx(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    # 表格文本也抽取
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)

def extract_txt(path):
    for enc in ("utf-8", "gbk", "gb18030", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except Exception:
            continue
    return ""

def extract_text(path):
    ext = Path(path).suffix.lower()
    try:
        if ext == ".pdf":
            return extract_pdf(path)
        if ext == ".docx":
            return extract_docx(path)
        if ext in (".txt", ".md"):
            return extract_txt(path)
        if ext == ".doc":
            # 旧版 .doc 需要 antiword/libreoffice；这里尝试 textract，失败则跳过
            try:
                import textract
                return textract.process(path).decode("utf-8", "ignore")
            except Exception:
                print(f"[跳过] 无法解析旧版 .doc（建议先转为 .docx）：{os.path.basename(path)}")
                return ""
    except Exception as e:
        print(f"[错误] 解析失败 {os.path.basename(path)}: {e}")
        return ""
    return ""

def get_doc_title(text, max_lines=5):
    """提取“文档标题区”：正文开头的前若干非空行（政策文件的大标题通常在最前面，
    有时第一行是文号，故合并前几行一起作为标题区，识别更稳）。"""
    lines = [l.strip() for l in (text or "").splitlines() if l.strip()]
    return " ".join(lines[:max_lines])[:200]

def detect_province(filename, text):
    """识别省份归属：优先用“PDF标题”（文件名 + 文档开头大标题），其次正文前部，
    再次按国家级文件兜底，仍无则返回‘未识别’。"""
    title_zone = f"{filename or ''} {get_doc_title(text)}"   # 标题区（最高优先级）
    body_zone = (text or "")[:3000]                          # 正文前部（次优先级）

    def best_match(zone):
        hits = {}
        for prov, kws in PROVINCE_KEYWORDS.items():
            score = sum(zone.count(kw) for kw in kws)
            if score:
                hits[prov] = score
        return max(hits, key=hits.get) if hits else None

    # 1) 先从标题区识别（文件名/文档大标题里通常直接写明“XX省/市”）
    prov = best_match(title_zone)
    if prov:
        return prov
    # 2) 标题没有，再看正文前部
    prov = best_match(body_zone)
    if prov:
        return prov
    # 3) 国家级文件兜底
    if any(k in title_zone or k in body_zone for k in NATIONAL_KEYWORDS):
        return "国家级"
    return "未识别"

# %%
# ====== 2.1 收集并抽取所有政策文本 ======
PATTERNS = ["*.pdf", "*.docx", "*.doc", "*.txt", "*.md"]
files = []
for pat in PATTERNS:
    files.extend(glob.glob(os.path.join(POLICY_DIR, "**", pat), recursive=True))
files = sorted(set(files))
print(f"在 {POLICY_DIR} 共发现 {len(files)} 个候选文件")

documents = {}   # filename -> {"path","text","province","n_chars"}
for fp in files:
    name = os.path.basename(fp)
    txt = extract_text(fp)
    if len(txt.strip()) < 80:   # 过滤空文件/解析失败
        continue
    documents[name] = {
        "path": fp,
        "text": txt,
        "province": detect_province(name, txt),
        "n_chars": len(txt),
    }
print(f"成功抽取有效文本的文件数：{len(documents)}")
for n, d in documents.items():
    print(f"  - {n}  | 省份={d['province']} | 字数={d['n_chars']}")

# %% [markdown]
# ## 4. DeepSeek 三级编码（开放→主轴→核心）+ 双维度标签
#
# 采用扎根理论式三级编码：**一级编码（初始概念）→ 二级编码（主轴范畴）→ 三级编码（核心范畴=政策工具）**；
# 同时为每条措施打上正交的**业务维度（产业链环节）**标签，形成“业务×工具”双维度。
# 每条编码必须附政策原文逐字引证，判定依据须结合政策工具理论。长文本自动分块合并、结果缓存。

# %%
client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)

def load_cache():
    if os.path.exists(CACHE_PATH):
        with open(CACHE_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_cache(cache):
    with open(CACHE_PATH, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)

def chunk_text(text, size=9000, overlap=400):
    """按字符分块，保留少量重叠以免切断语义。"""
    if len(text) <= size:
        return [text]
    chunks, i = [], 0
    while i < len(text):
        chunks.append(text[i:i + size])
        i += size - overlap
    return chunks

def call_deepseek_json(system_prompt, user_prompt, max_retries=4, temperature=0.2):
    """调用 DeepSeek 并强制 JSON 输出，带指数退避重试。"""
    for attempt in range(max_retries):
        try:
            resp = client.chat.completions.create(
                model=DEEPSEEK_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=temperature,
                response_format={"type": "json_object"},
            )
            return json.loads(resp.choices[0].message.content)
        except Exception as e:
            wait = 2 ** attempt
            print(f"  [重试 {attempt+1}/{max_retries}] {e}，{wait}s 后重试…")
            time.sleep(wait)
    raise RuntimeError("DeepSeek 调用多次失败")

# %%
CODING_SYSTEM = (
    "你是一名严谨的卫生政策质性研究编码员，精通扎根理论三级编码与政策工具理论"
    "（Rothwell & Zegveld 供给型/环境型/需求型）。你只依据给定文本编码，绝不臆测或补充文本中不存在的信息。"
    "一级编码须概念化、精炼；判定政策工具时须结合工具的理论定义进行推断，而非仅做字面关键词匹配。"
    "每条编码都必须能在原文中找到支撑句并逐字引证。"
)

def build_coding_prompt(doc_name, province, chunk_text_):
    biz = "、".join(BUSINESS_DIMENSIONS)
    tools = "；".join(f"{t}（{POLICY_TOOLS[t]['定义']}）" for t in TOOL_TYPES)
    subs = "、".join(TOOL_SUBTYPES)
    return f"""请对下面这份政策文本片段做质性编码（扎根理论式三级编码）。文件：《{doc_name}》（疑似省份：{province}）。
逐条抽取文本中**明确表达的具体政策措施**，每条给出三级编码与双维度标签。

【编码规范】
1. 一级编码（初始概念）：对该措施的概念化命名，**精炼名词短语（建议≤12字）**，
   例如“攻关平台建设”“中药二次开发奖励”“审评审批绿色通道”；**禁止照抄长句或写成动作摘要**。
2. 二级编码（主轴范畴）：把语义相近的一级概念归并的中层范畴名，如“平台与基础设施”“资金支持与激励”“审评审批优化”“市场准入与推广”。
3. 业务维度（产业链/创新链环节，X轴，单选）：{biz}；确属跨环节统筹填“全链条统筹”。
4. 三级编码（核心范畴＝政策工具，Y轴，单选）：{tools}
5. 工具细类（单选，参考 McDonnell&Elmore）：{subs}
6. 药物类型（单选）：{" / ".join(DRUG_TYPES)}（无法区分填“通用”）。
7. 原文：逐字引证支撑句（可截取关键句）。
8. 判定依据：**结合上述政策工具的理论定义**说明为何归入该三级工具（须体现理论判定，不要只罗列关键词）。

严格输出 JSON：
{{"编码": [
  {{"一级编码": "...", "二级编码": "...", "业务维度": "...", "三级编码": "...", "工具细类": "...",
    "药物类型": "...", "原文": "...", "判定依据": "..."}}
]}}
只编码文本明确表达的内容，不得臆造。若本片段无实质措施，返回 {{"编码": []}}。

【待编码文本片段】
\"\"\"{chunk_text_}\"\"\"
"""

def merge_codings(codings):
    """合并同一文件多个分块的编码；按‘一级编码’去重。"""
    merged = {"编码": []}
    seen = set()
    for c in codings:
        for item in c.get("编码", []) or []:
            k = (item.get("一级编码") or "").strip()
            if not k or k in seen:
                continue
            seen.add(k)
            merged["编码"].append(item)
    return merged

def code_document(doc_name, doc, cache, force=False):
    if (doc_name in cache) and not force:
        return cache[doc_name]
    chunks = chunk_text(doc["text"])
    results = []
    for idx, ch in enumerate(chunks):
        print(f"    编码分块 {idx+1}/{len(chunks)} …")
        prompt = build_coding_prompt(doc_name, doc["province"], ch)
        results.append(call_deepseek_json(CODING_SYSTEM, prompt))
    merged = merge_codings(results)
    merged["_province"] = doc["province"]
    cache[doc_name] = merged
    save_cache(cache)
    return merged

# %%
# ====== 4.1 执行编码（带缓存）======
cache = load_cache()
for name, doc in tqdm(documents.items(), desc="编码文件"):
    if name in cache:
        continue
    print(f"\n>>> 正在编码：{name}")
    try:
        code_document(name, doc, cache)
    except Exception as e:
        print(f"  [失败] {name}: {e}")
print("\n编码完成。已缓存到:", CACHE_PATH)

# %% [markdown]
# ## 5. 质性编码汇总表（三级编码 + 双维度）
# 生成长表：文件 × 一级/二级/三级编码 × 业务维度 × 药物类型 × 原文引证 × 判定依据，并导出附录A。

# %%
rows = []
for name, doc in documents.items():
    coded = cache.get(name)
    if not coded:
        continue
    prov = doc["province"]
    for item in coded.get("编码", []) or []:
        biz = (item.get("业务维度", "") or "").strip()
        biz = biz if biz in BUSINESS_DIMENSIONS else "全链条统筹"
        tool = (item.get("三级编码", "") or "").strip()
        if tool not in TOOL_TYPES:   # LLM 未给合法工具时用关键词兜底
            tool = classify_tool_by_keyword(f"{item.get('一级编码','')} {item.get('原文','')}")
        yiji = item.get("一级编码", "")
        rows.append({
            "文件": name, "省份": prov,
            "一级编码": yiji,
            "二级编码": item.get("二级编码", ""),
            "三级编码": tool,
            "工具细类": item.get("工具细类", ""),
            "业务维度": biz,
            "药物类型": item.get("药物类型", "通用"),
            "原文引证": item.get("原文", ""),
            "判定依据": item.get("判定依据", ""),
            # —— 兼容别名（供下游沿用旧字段名）——
            "编码要点": yiji,
            "维度": biz,
            "工具类型": tool,
        })

coding_df = pd.DataFrame(rows)
print("编码总条数:", len(coding_df))

# 导出附录A（体现三级编码层级 + 双维度，规范表头）
if not coding_df.empty:
    appendixA = coding_df[["文件", "省份", "一级编码", "二级编码", "三级编码", "工具细类",
                           "业务维度", "药物类型", "判定依据", "原文引证"]].rename(columns={
        "一级编码": "一级编码(初始概念)", "二级编码": "二级编码(主轴范畴)",
        "三级编码": "三级编码(核心范畴·政策工具)", "工具细类": "工具细类(五类)",
        "业务维度": "业务维度(产业链环节)", "判定依据": "判定依据(理论)"})
    coding_xlsx = os.path.join(OUTPUT_DIR, "附录A_质性编码汇总表.xlsx")
    appendixA.to_excel(coding_xlsx, index=False)
    print("已导出（三级编码 + 双维度）:", coding_xlsx)
coding_df.head(20)

# %% [markdown]
# ## 6. 政策工具维度分析（Rothwell & Zegveld：供给型/环境型/需求型）
#
# 政策工具（三级核心范畴）已在第 4 节编码时**逐条结合理论判定**（判定依据见附录A），
# 本节据此统计各类工具的数量、占比与结构性趋势，并做“工具 × 业务维度/省份/药物类型”交叉。

# %%
# ====== 6.1 表2：政策工具类型（Rothwell）分布及占比 ======
tool_counts = coding_df["工具类型"].value_counts().reindex(TOOL_TYPES, fill_value=0)
tool_total = int(tool_counts.sum()) or 1
tool_table = pd.DataFrame({
    "政策工具(三级核心范畴)": TOOL_TYPES,
    "理论定义": [POLICY_TOOLS[t]["定义"] for t in TOOL_TYPES],
    "政策条目数": [int(tool_counts[t]) for t in TOOL_TYPES],
    "占比(%)": [round(tool_counts[t] / tool_total * 100, 1) for t in TOOL_TYPES],
    "代表性一级编码": ["；".join(coding_df[coding_df["工具类型"] == t]["编码要点"].head(3).tolist()) for t in TOOL_TYPES],
})
tool_table_xlsx = os.path.join(OUTPUT_DIR, "表2_政策工具类型分布.xlsx")
tool_table.to_excel(tool_table_xlsx, index=False)
print("已导出:", tool_table_xlsx)
tool_table

# %%
# ====== 6.2 工具类型 × 省份 / 业务维度 / 药物类型 交叉表 ======
tool_by_prov = (coding_df.pivot_table(index="省份", columns="工具类型", values="编码要点",
                                       aggfunc="count", fill_value=0)
                .reindex(columns=TOOL_TYPES, fill_value=0))
tool_by_dim = (coding_df[coding_df["维度"].isin(DIMENSIONS)]
               .pivot_table(index="维度", columns="工具类型", values="编码要点", aggfunc="count", fill_value=0)
               .reindex(index=DIMENSIONS, columns=TOOL_TYPES, fill_value=0))
tool_by_drug = (coding_df.pivot_table(index="药物类型", columns="工具类型", values="编码要点",
                                      aggfunc="count", fill_value=0)
                .reindex(columns=TOOL_TYPES, fill_value=0))
with pd.ExcelWriter(os.path.join(OUTPUT_DIR, "表2附_工具类型交叉表.xlsx")) as xw:
    tool_by_prov.to_excel(xw, sheet_name="工具x省份")
    tool_by_dim.to_excel(xw, sheet_name="工具x维度")
    tool_by_drug.to_excel(xw, sheet_name="工具x药物类型")
print("已导出工具类型交叉表")
tool_by_dim

# %% [markdown]
# ## 7. 归纳式主题分析（跨维度）+ 两种方法的整合
#
# **方法定位（避免与第3节维度比较重复）**：
# - 第 3 节的“质性分析”是**框架式定向内容分析**（演绎）：沿 7 个先验维度做结构化描述与横向比较。
# - 本节“主题分析”是 Braun & Clarke 的**反身性主题分析**（归纳）：把**全部维度的编码混合在一起**，
#   提炼能够**横跨多个维度**的潜在主题（latent themes），捕捉张力与趋势——主题不再按维度切分，
#   因此不会与第 3 节重复。
# - 7.2 给出**整合（联合展示 joint display）**：主题 × 维度 的对应矩阵，做三角互证，体现两种方法的互补。

# %%
# ====== 7.1 跨维度归纳主题分析 ======
THEME_SYSTEM = (
    "你是质性研究主题分析专家，精通 Braun & Clarke 的反身性主题分析（reflexive thematic analysis）。"
    "你要在所有编码之上做归纳式提炼，生成能够横跨多个政策维度的潜在主题（latent theme），"
    "而不是简单复述给定的维度名称。只能基于给定编码归纳，不得引入外部信息或编造省份/政策。"
)

def build_theme_prompt(all_items):
    # 去重并附带维度/省份/工具类型标签，供模型做跨维度归纳
    seen, payload = set(), []
    for it in all_items:
        code = it["编码要点"].strip()
        if not code or code in seen:
            continue
        seen.add(code)
        payload.append({"编码": code, "维度": it["维度"], "省份": it["省份"], "工具": it.get("工具类型", "")})
    return f"""下面是来自多份省级创新医药政策文本、覆盖全部评估维度的编码要点（JSON 数组，已含其所属维度标签）。
请做**跨维度的归纳式主题分析**：在所有编码之上提炼 4-7 个**潜在核心主题**。

关键要求：
1. 主题应**横跨/整合多个维度**，揭示深层模式、张力或趋势，**不要**简单照搬“创新要素配置”等维度名作为主题。
2. 每个主题给出：主题名称、主题阐释（2-4句，点明它如何贯穿不同维度）、涉及维度（可多个）、
   覆盖省份、支持文件数、3-5 条代表性编码（须从输入中原样挑选，注明其维度）。
3. 忠实于输入编码，禁止臆造。
4. 严格输出 JSON：
{{"主题": [{{"主题名称": "...", "主题阐释": "...", "涉及维度": ["..."], "覆盖省份": ["..."], "支持文件数": 0,
            "代表编码": [{{"编码": "...", "维度": "..."}}]}}]}}

【输入编码】
{json.dumps(payload, ensure_ascii=False)}
"""

_all_dim_items = [r for r in rows if r["维度"] in DIMENSIONS]
if len(_all_dim_items) >= 3:
    theme_results = call_deepseek_json(THEME_SYSTEM, build_theme_prompt(_all_dim_items))
else:
    theme_results = {"主题": []}

with open(os.path.join(OUTPUT_DIR, "主题分析结果.json"), "w", encoding="utf-8") as f:
    json.dump(theme_results, f, ensure_ascii=False, indent=2)
print(f"跨维度主题分析完成，共提炼主题 {len(theme_results.get('主题', []))} 个。")

# 主题分析表
theme_rows = []
for t in theme_results.get("主题", []) or []:
    reps = t.get("代表编码", []) or []
    rep_str = " || ".join(r.get("编码", "") if isinstance(r, dict) else str(r) for r in reps)
    theme_rows.append({
        "主题名称": t.get("主题名称", ""),
        "主题阐释": t.get("主题阐释", ""),
        "涉及维度": "、".join(t.get("涉及维度", []) or []),
        "覆盖省份": "、".join(t.get("覆盖省份", []) or []),
        "支持文件数": t.get("支持文件数", 0),
        "代表编码": rep_str,
    })
theme_df = pd.DataFrame(theme_rows)
theme_xlsx = os.path.join(OUTPUT_DIR, "主题分析表.xlsx")
theme_df.to_excel(theme_xlsx, index=False)
print("已导出:", theme_xlsx)
theme_df

# %%
# ====== 7.2 整合分析：主题 × 维度 联合展示（joint display，三角互证）======
# 每个主题在各维度上的“覆盖”：以模型给出的“涉及维度”为主，并用代表编码回链编码表做证据校验。
def theme_dimension_vector(theme):
    dims = set(theme.get("涉及维度", []) or [])
    for rep in theme.get("代表编码", []) or []:
        code = rep.get("编码", "") if isinstance(rep, dict) else str(rep)
        d = rep.get("维度", "") if isinstance(rep, dict) else ""
        if d in DIMENSIONS:
            dims.add(d)
        if code:
            hit = coding_df[coding_df["编码要点"].apply(lambda x: bool(x) and (x in code or code in x))]
            dims.update([x for x in hit["维度"].tolist() if x in DIMENSIONS])
    return {d: (1 if d in dims else 0) for d in DIMENSIONS}

joint_rows = []
for t in theme_results.get("主题", []) or []:
    vec = theme_dimension_vector(t)
    row = {"主题": t.get("主题名称", "")}
    row.update({d: vec[d] for d in DIMENSIONS})
    row["跨越维度数"] = sum(vec.values())
    joint_rows.append(row)
joint_df = pd.DataFrame(joint_rows).set_index("主题") if joint_rows else pd.DataFrame(columns=["主题"] + DIMENSIONS).set_index("主题")
joint_xlsx = os.path.join(OUTPUT_DIR, "表_整合联合展示_主题x维度.xlsx")
joint_df.to_excel(joint_xlsx)
print("已导出整合联合展示表:", joint_xlsx)
print("说明：1=该主题贯穿此维度。‘跨越维度数’越大，说明该主题越是跨维度的整合性主题。")
joint_df

# %%
# ====== 7.3 省份 × 主题 覆盖矩阵（核心主题的空间分布）======
tp_rows = []
for t in theme_results.get("主题", []) or []:
    covered = set(t.get("覆盖省份", []) or [])
    # 用代表编码回链编码表，补充省份证据（防止仅凭模型声明）
    for rep in t.get("代表编码", []) or []:
        code = rep.get("编码", "") if isinstance(rep, dict) else str(rep)
        if code:
            hit = coding_df[coding_df["编码要点"].apply(lambda x: bool(x) and (x in code or code in x))]
            covered.update(hit["省份"].tolist())
    for p in covered:
        if p:
            tp_rows.append({"省份": p, "主题": t.get("主题名称", "")})
if tp_rows:
    theme_by_prov = (pd.DataFrame(tp_rows).assign(v=1)
                     .pivot_table(index="省份", columns="主题", values="v", aggfunc="max", fill_value=0))
else:
    theme_by_prov = pd.DataFrame()
tp_xlsx = os.path.join(OUTPUT_DIR, "表_省份x主题_覆盖.xlsx")
theme_by_prov.to_excel(tp_xlsx)
print("已导出省份×主题覆盖表:", tp_xlsx)
theme_by_prov

# %% [markdown]
# ## 8. 比较矩阵
# （1）省份 × 维度 编码数量矩阵；（2）化学药/生物制品 vs 中医药 对照表（含参考列）。

# %%
# ====== 8.1 省份 × 维度 编码数量矩阵 ======
prov_dim = (coding_df[coding_df["维度"].isin(DIMENSIONS)]
            .pivot_table(index="省份", columns="维度", values="编码要点",
                         aggfunc="count", fill_value=0))
# 维度列排序
prov_dim = prov_dim.reindex(columns=DIMENSIONS, fill_value=0)
matrix_xlsx = os.path.join(OUTPUT_DIR, "省份_维度_编码矩阵.xlsx")
prov_dim.to_excel(matrix_xlsx)
print("已导出:", matrix_xlsx)
prov_dim

# %%
# ====== 8.2 化学药/生物制品 vs 中医药 对照表 ======
def summarize_codes(dim, drug_type, max_items=6):
    sub = coding_df[(coding_df["维度"] == dim) & (coding_df["药物类型"] == drug_type)]
    pts = list(dict.fromkeys(sub["编码要点"].tolist()))  # 去重保序
    return "；".join(pts[:max_items]) if pts else "（现有政策文本未涉及）"

cmp_rows = []
for d in DIMENSIONS:
    cmp_rows.append({
        "评估维度": d,
        "参考-化学药与生物制品": REFERENCE_MATRIX.get(d, {}).get("化学药与生物制品", ""),
        "参考-中医药": REFERENCE_MATRIX.get(d, {}).get("中医药", ""),
        "实证-化学药与生物制品": summarize_codes(d, "化学药与生物制品"),
        "实证-中医药": summarize_codes(d, "中医药"),
        "实证-通用": summarize_codes(d, "通用"),
    })
compare_df = pd.DataFrame(cmp_rows)
compare_xlsx = os.path.join(OUTPUT_DIR, "化学药生物制品_vs_中医药_对照表.xlsx")
compare_df.to_excel(compare_xlsx, index=False)
print("已导出:", compare_xlsx)
compare_df

# %%
# ====== 8.3 政策工具层面（Rothwell）化药/生物制品 vs 中医药 对照 ======
def summarize_codes_by_tool(tool, drug_type, max_items=6):
    sub = coding_df[(coding_df["工具类型"] == tool) & (coding_df["药物类型"] == drug_type)]
    pts = list(dict.fromkeys(sub["编码要点"].tolist()))
    return "；".join(pts[:max_items]) if pts else "（现有政策文本未涉及）"

tool_cmp_rows = []
for t in TOOL_TYPES:
    tool_cmp_rows.append({
        "政策工具": t,
        "参考-化学药与生物制品": REFERENCE_TOOL.get(t, {}).get("化学药与生物制品", ""),
        "参考-中医药": REFERENCE_TOOL.get(t, {}).get("中医药", ""),
        "实证-化学药与生物制品": summarize_codes_by_tool(t, "化学药与生物制品"),
        "实证-中医药": summarize_codes_by_tool(t, "中医药"),
        "实证-通用": summarize_codes_by_tool(t, "通用"),
    })
tool_compare_df = pd.DataFrame(tool_cmp_rows)
tool_compare_xlsx = os.path.join(OUTPUT_DIR, "化药生物制品_vs_中医药_政策工具对照表.xlsx")
tool_compare_df.to_excel(tool_compare_xlsx, index=False)
print("已导出:", tool_compare_xlsx)
tool_compare_df

# %% [markdown]
# ## 9. 可视化
# 中文字体：脚本会尝试常见中文字体；若图中中文显示为方框，请安装并指定 SimHei/Microsoft YaHei。

# %%
# 中文字体自适应
def set_chinese_font():
    candidates = ["Microsoft YaHei", "SimHei", "PingFang SC", "Heiti SC",
                  "WenQuanYi Zen Hei", "Noto Sans CJK SC", "Source Han Sans SC", "Arial Unicode MS"]
    available = {f.name for f in font_manager.fontManager.ttflist}
    for c in candidates:
        if c in available:
            matplotlib.rcParams["font.sans-serif"] = [c]
            matplotlib.rcParams["axes.unicode_minus"] = False
            print("使用中文字体:", c)
            return
    print("[警告] 未找到中文字体，图中中文可能显示为方框。请安装 SimHei 等字体。")
    matplotlib.rcParams["axes.unicode_minus"] = False
set_chinese_font()

FIG_DIR = os.path.join(OUTPUT_DIR, "图")
os.makedirs(FIG_DIR, exist_ok=True)

# %%
# ====== 图1：省份 × 维度 编码数量热力图 ======
if not prov_dim.empty:
    fig, ax = plt.subplots(figsize=(11, max(4, 0.6 * len(prov_dim) + 2)))
    data = prov_dim.values
    im = ax.imshow(data, cmap="YlOrRd", aspect="auto")
    ax.set_xticks(range(len(prov_dim.columns)))
    ax.set_xticklabels(prov_dim.columns, rotation=30, ha="right")
    ax.set_yticks(range(len(prov_dim.index)))
    ax.set_yticklabels(prov_dim.index)
    for i in range(data.shape[0]):
        for j in range(data.shape[1]):
            ax.text(j, i, int(data[i, j]), ha="center", va="center", fontsize=9)
    ax.set_title("各省份在不同评估维度上的政策编码数量分布")
    fig.colorbar(im, ax=ax, label="编码条数")
    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, "图1_省份维度热力图.png"), dpi=200)
    plt.show()

# %%
# ====== 图2：各维度编码总数条形图 ======
dim_counts = (coding_df[coding_df["维度"].isin(DIMENSIONS)]
              .groupby("维度").size().reindex(DIMENSIONS, fill_value=0))
fig, ax = plt.subplots(figsize=(10, 5))
ax.bar(dim_counts.index, dim_counts.values, color="#4C72B0")
ax.set_title("各评估维度的政策编码总数")
ax.set_ylabel("编码条数")
plt.xticks(rotation=25, ha="right")
for i, v in enumerate(dim_counts.values):
    ax.text(i, v, str(int(v)), ha="center", va="bottom")
plt.tight_layout()
plt.savefig(os.path.join(FIG_DIR, "图2_维度编码总数.png"), dpi=200)
plt.show()

# %%
# ====== 图3：化学药/生物制品 vs 中医药 各维度编码对比 ======
pivot_drug = (coding_df[coding_df["维度"].isin(DIMENSIONS)]
              .pivot_table(index="维度", columns="药物类型", values="编码要点",
                           aggfunc="count", fill_value=0).reindex(DIMENSIONS, fill_value=0))
for col in DRUG_TYPES:
    if col not in pivot_drug.columns:
        pivot_drug[col] = 0
pivot_drug = pivot_drug[DRUG_TYPES]
fig, ax = plt.subplots(figsize=(11, 5))
x = np.arange(len(pivot_drug.index)); w = 0.27
colors = {"化学药与生物制品": "#C44E52", "中医药": "#55A868", "通用": "#8172B3"}
for k, col in enumerate(DRUG_TYPES):
    ax.bar(x + (k - 1) * w, pivot_drug[col].values, w, label=col, color=colors[col])
ax.set_xticks(x); ax.set_xticklabels(pivot_drug.index, rotation=25, ha="right")
ax.set_ylabel("编码条数"); ax.set_title("化学药与生物制品 vs 中医药：各维度政策着力点对比")
ax.legend()
plt.tight_layout()
plt.savefig(os.path.join(FIG_DIR, "图3_化药生物制品_vs_中医药.png"), dpi=200)
plt.show()

# %%
# ====== 图4：核心主题支持强度（按支持文件数）======
if not theme_df.empty:
    td = theme_df.copy()
    td["标签"] = td["主题名称"]
    td = td.sort_values("支持文件数", ascending=True).tail(15)
    fig, ax = plt.subplots(figsize=(11, max(4, 0.5 * len(td) + 2)))
    ax.barh(td["标签"], td["支持文件数"], color="#DD8452")
    ax.set_xlabel("支持文件数")
    ax.set_title("跨维度核心主题的支持强度")
    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, "图4_主题支持强度.png"), dpi=200)
    plt.show()

# %%
# ====== 图5：政策工具类型分布（表2）======
fig, ax = plt.subplots(figsize=(9, 5))
_vals = [int(tool_counts[t]) for t in TOOL_TYPES]
_palette = ["#C44E52", "#DD8452", "#55A868", "#4C72B0", "#8172B3"]
bars = ax.bar(TOOL_TYPES, _vals, color=_palette)
ax.set_title("政策工具类型分布（表2）"); ax.set_ylabel("政策条目数")
_tot = sum(_vals) or 1
for b, v in zip(bars, _vals):
    ax.text(b.get_x() + b.get_width() / 2, v, f"{v}\n{round(v/_tot*100,1)}%",
            ha="center", va="bottom", fontsize=9)
plt.xticks(rotation=20, ha="right"); plt.tight_layout()
plt.savefig(os.path.join(FIG_DIR, "图5_政策工具类型分布.png"), dpi=200)
plt.show()

# %%
# ====== 图6：工具类型 × 评估维度 构成（堆叠条形）======
fig, ax = plt.subplots(figsize=(11, 6))
bottom = np.zeros(len(tool_by_dim.index))
for k, t in enumerate(TOOL_TYPES):
    ax.bar(tool_by_dim.index, tool_by_dim[t].values, bottom=bottom, label=t, color=_palette[k % 5])
    bottom += tool_by_dim[t].values
ax.set_title("各评估维度下的政策工具类型构成"); ax.set_ylabel("政策条目数")
ax.legend()
plt.xticks(rotation=25, ha="right"); plt.tight_layout()
plt.savefig(os.path.join(FIG_DIR, "图6_工具类型x维度.png"), dpi=200)
plt.show()

# %%
# ====== 图7：整合联合展示——主题 × 维度 热力图（体现两种方法的整合）======
if not joint_df.empty:
    jm = joint_df[DIMENSIONS]
    fig, ax = plt.subplots(figsize=(11, max(4, 0.6 * len(jm) + 2)))
    im = ax.imshow(jm.values, cmap="Greens", aspect="auto", vmin=0, vmax=1)
    ax.set_xticks(range(len(DIMENSIONS))); ax.set_xticklabels(DIMENSIONS, rotation=30, ha="right")
    ax.set_yticks(range(len(jm.index))); ax.set_yticklabels(jm.index)
    for i in range(jm.shape[0]):
        for j in range(jm.shape[1]):
            if jm.values[i, j]:
                ax.text(j, i, "✓", ha="center", va="center", fontsize=11, color="#114411")
    ax.set_title("整合联合展示：归纳主题如何贯穿先验维度（✓=该主题覆盖此维度）")
    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, "图7_整合联合展示.png"), dpi=200)
    plt.show()

# %%
# ====== 图8：省份 × 核心主题 覆盖热力图（主题的空间分布）======
if not theme_by_prov.empty:
    m = theme_by_prov
    fig, ax = plt.subplots(figsize=(max(8, 1.1 * len(m.columns) + 3), max(4, 0.5 * len(m.index) + 2)))
    im = ax.imshow(m.values, cmap="Blues", aspect="auto", vmin=0, vmax=1)
    ax.set_xticks(range(len(m.columns))); ax.set_xticklabels(m.columns, rotation=30, ha="right")
    ax.set_yticks(range(len(m.index))); ax.set_yticklabels(m.index)
    for i in range(m.shape[0]):
        for j in range(m.shape[1]):
            if m.values[i, j]:
                ax.text(j, i, "●", ha="center", va="center", color="#1f4e79", fontsize=11)
    ax.set_title("省份 × 核心主题 覆盖分布（●=该省政策支持此主题）")
    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, "图8_省份x主题覆盖.png"), dpi=200)
    plt.show()

print("图表已保存到:", FIG_DIR)

# %% [markdown]
# ## 10. 生成约 8000 字研究报告
#
# 仅向模型提供**已抽取的真实编码、主题与矩阵**作为证据，分章节生成以保证篇幅与质量。
# 强制要求：凡证据不足处标注“现有政策文本未涉及”，引用观点须标注来源文件/省份，禁止杜撰。

# %%
# ====== 10.1 构建“证据包”（提供给模型的唯一事实来源）======
def build_evidence_pack():
    lines = []
    lines.append("【一、纳入文本清单】")
    for name, doc in documents.items():
        lines.append(f"- 《{name}》 省份={doc['province']} 字数={doc['n_chars']}")

    lines.append("\n【二、省份×维度 编码数量矩阵】")
    lines.append(prov_dim.to_string())

    lines.append("\n【三、各维度编码要点（含省份与药物类型；这是唯一可引用的事实证据）】")
    for d in DIMENSIONS:
        lines.append(f"\n## 维度：{d}")
        sub = [r for r in rows if r["维度"] == d]
        for r in sub:
            q = (r["原文引证"] or "")[:120]
            lines.append(f"  - [{r['省份']}|{r['药物类型']}] {r['编码要点']}（原文：{q}）")

    lines.append("\n【四、归纳式跨维度主题分析结果（解释层，独立于第3节的维度描述）】")
    for t in theme_results.get("主题", []) or []:
        lines.append(f"  - 主题：{t.get('主题名称','')}；涉及维度：{'、'.join(t.get('涉及维度',[]) or [])}；"
                     f"省份：{'、'.join(t.get('覆盖省份',[]) or [])}；支持文件数：{t.get('支持文件数',0)}；"
                     f"阐释：{t.get('主题阐释','')}")
    if not joint_df.empty:
        lines.append("\n【四附、整合联合展示：主题 × 维度（1=该主题贯穿此维度，体现两方法互补）】")
        lines.append(joint_df.to_string())
    if not theme_by_prov.empty:
        lines.append("\n【四附2、省份 × 核心主题 覆盖（1=该省政策支持此主题）】")
        lines.append(theme_by_prov.to_string())

    lines.append("\n【五、化学药/生物制品 vs 中医药 实证对照（来自编码）】")
    for _, r in compare_df.iterrows():
        lines.append(f"  - {r['评估维度']}：化学药/生物制品=[{r['实证-化学药与生物制品']}]；"
                     f"中医药=[{r['实证-中医药']}]；通用=[{r['实证-通用']}]")

    lines.append("\n【六、省份产业画像（背景参考资料，非政策文本证据，引用须注明为背景）】")
    for p, prof in PROVINCE_PROFILE.items():
        lines.append(f"  - {p}：{prof}")

    lines.append("\n【七、政策工具类型分布（表2；Rothwell & Zegveld：供给型/环境型/需求型）】")
    for t in TOOL_TYPES:
        lines.append(f"  - {t}：{int(tool_counts[t])} 条，占 {round(tool_counts[t]/tool_total*100,1)}%")
    lines.append("工具类型 × 业务维度 构成（行=业务维度，列=工具类型）：")
    lines.append(tool_by_dim.to_string())
    lines.append("工具类型 × 省份 构成：")
    lines.append(tool_by_prov.to_string())
    lines.append("\n【七附、政策工具层面 化药/生物制品 vs 中医药 对照（Rothwell）】")
    for _, r in tool_compare_df.iterrows():
        lines.append(f"  - {r['政策工具']}：化药/生物制品=[{r['实证-化学药与生物制品']}]；中医药=[{r['实证-中医药']}]")
    return "\n".join(lines)

EVIDENCE = build_evidence_pack()
print("证据包字符数:", len(EVIDENCE))
print(EVIDENCE[:1500])

# %%
REPORT_SYSTEM = (
    "你是一名资深卫生政策研究学者，撰写规范的中文学术研究报告。"
    "你必须严格基于用户提供的【证据包】写作：所有事实性陈述都要能在证据包中找到依据，"
    "引用具体政策要点时标注来源省份/文件；证据不足之处必须写明‘现有政策文本未涉及’，"
    "严禁编造政策条款、数字、省份或文件。语言客观、学术、有逻辑层次。"
)

# 分章节生成（整合式混合方法设计：框架式质性分析 + 归纳式主题分析 + 整合）
REPORT_SECTIONS = [
    ("摘要与关键词", "撰写中文摘要（研究目的/方法/主要发现/价值，约400字）与5个关键词。点明采用‘框架式质性分析+归纳式主题分析’的整合设计。", 450),
    ("一、引言", "问题提出、研究背景（可用省份产业画像作背景）、研究意义与研究问题。", 900),
    ("二、研究方法", "阐明整合式混合方法设计：①框架式质性分析（演绎）——采用扎根理论三级编码（一级初始概念→二级主轴范畴→三级核心范畴），沿产业链‘业务维度’与‘政策工具维度’（Rothwell&Zegveld：供给型/环境型/需求型）构成双维度矩阵；②Braun&Clarke反身性主题分析（归纳，跨维度提炼潜在主题）；③两者如何整合（联合展示joint display、三角互证、互补而非重复）。说明DeepSeek辅助三级编码流程、文本遴选、防杜撰与效度保障（逐字引证、理论化判定依据、缓存可复现）。", 1300),
    ("三、框架式质性分析：业务维度（产业链环节）横向比较", "这是‘结构性/描述性’发现层。沿6个业务维度（创新要素配置、创新方向、临床前研究、临床研究、成果转化、生产与市场准入）做省份差异比较，并对照化学药/生物制品与中医药两条路径。多引用证据包中的具体一级编码与省份、给出频次。本章只做结构化描述，不做跨维度归纳。", 2000),
    ("四、归纳式主题分析：跨维度核心主题", "这是‘解释性’发现层，必须与第三章区分开。呈现归纳出的4-7个**跨维度潜在主题**（来自证据包【四】），重点阐述每个主题如何横跨多个维度、揭示了何种深层模式或张力，**不要按维度复述、不要重复第三章**。", 1100),
    ("五、两种方法的整合：联合展示与三角互证", "基于证据包【四附】的‘主题×维度’联合展示矩阵，说明归纳主题与先验维度如何对应整合：哪些主题跨越多个维度（整合性强）、两种方法在何处相互印证（三角互证收敛）、在何处各有侧重（互补）。讨论这种整合设计对结论稳健性的意义。", 950),
    ("六、政策工具类型结构与趋势", "基于证据包【七】表2数据，分析三类政策工具（Rothwell&Zegveld：供给型/环境型/需求型）的数量与占比结构，结合‘工具×业务维度/省份’交叉揭示工具偏好与结构性趋势（须以实际数字为准），并结合【七附】对照讨论化药与中医药在工具运用上的差异及其治理逻辑与不足。", 1000),
    ("七、省域格局与化药/中医药路径差异讨论", "讨论区域格局、化学药生物制品与中医药政策逻辑差异及其成因。", 900),
    ("八、政策建议与研究局限", "提出可操作的政策建议（含政策工具结构优化）；说明研究局限（样本范围、文本时效、编码与分类主观性、单一模型偏差等）与展望。", 900),
]

# 每章末尾内嵌的图表（标题关键词 -> [(图片文件, 图注)]）
FIGURE_MAP = {
    "三、框架式质性分析": [("图1_省份维度热力图.png", "图1 各省份在不同评估维度上的政策编码数量分布"),
                    ("图2_维度编码总数.png", "图2 各评估维度的政策编码总数"),
                    ("图3_化药生物制品_vs_中医药.png", "图3 化学药与生物制品 vs 中医药：各维度政策着力点对比")],
    "四、归纳式主题分析": [("图4_主题支持强度.png", "图4 跨维度核心主题的支持强度"),
                    ("图8_省份x主题覆盖.png", "图8 省份 × 核心主题 覆盖分布")],
    "五、两种方法的整合": [("图7_整合联合展示.png", "图7 整合联合展示：归纳主题如何贯穿先验维度")],
    "六、政策工具类型": [("图5_政策工具类型分布.png", "图5 政策工具类型分布"),
                  ("图6_工具类型x维度.png", "图6 各评估维度下的政策工具类型构成")],
}
def figures_for(title):
    for k, v in FIGURE_MAP.items():
        if title.startswith(k):
            return v
    return []

def generate_section(title, instruction, target_words, prev_titles):
    user = f"""【证据包】
{EVIDENCE}

【写作任务】
现在撰写研究报告的章节：「{title}」。
内容要求：{instruction}
目标字数：约 {target_words} 字（中文）。
已写章节（避免重复）：{prev_titles}

请直接输出该章节的正文（含本级标题），使用 Markdown。务必基于证据包，标注来源省份/文件，证据不足处标注‘现有政策文本未涉及’。"""
    resp = client.chat.completions.create(
        model=DEEPSEEK_MODEL,
        messages=[{"role": "system", "content": REPORT_SYSTEM},
                  {"role": "user", "content": user}],
        temperature=0.5,
    )
    return resp.choices[0].message.content

section_texts = {}   # 章节标题 -> 生成的正文（供 Word 导出复用，避免与图片混在一起）
report_parts = ["# 我国省域创新医药政策研究——框架式质性分析与归纳式主题分析的整合\n"]
prev = []
for title, instruction, tgt in REPORT_SECTIONS:
    print(f">>> 生成章节：{title}（目标约{tgt}字）")
    sec = generate_section(title, instruction, tgt, "；".join(prev))
    section_texts[title] = sec
    report_parts.append(sec)
    # 在该章正文后内嵌对应的图（Markdown 图片引用，相对路径指向“图/”目录）
    for fig_file, caption in figures_for(title):
        if os.path.exists(os.path.join(FIG_DIR, fig_file)):
            report_parts.append(f"\n![{caption}](图/{fig_file})\n\n*{caption}*\n")
    prev.append(title)
    time.sleep(1)

report_md = "\n\n".join(report_parts)
report_path = os.path.join(OUTPUT_DIR, "创新医药政策研究报告.md")
with open(report_path, "w", encoding="utf-8") as f:
    f.write(report_md)
approx_words = len(re.findall(r"[\u4e00-\u9fff]", report_md))
print(f"\nMarkdown 报告已生成（图表已内嵌）：{report_path}")
print(f"中文字数（汉字计）约：{approx_words}")

# %%
# ====== 10.3 导出内嵌图表的 Word 报告（.docx，适合投稿核心期刊）======
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _add_markdown_block(doc, text):
    """把一段 Markdown 文本粗略转为 Word：# 标题转标题样式，图片行跳过，其余为正文段落。"""
    for line in (text or "").splitlines():
        s = line.rstrip()
        if not s.strip() or s.lstrip().startswith("!["):
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", s.strip())
        if m:
            doc.add_heading(re.sub(r"[*#`]", "", m.group(2)).strip(), level=min(len(m.group(1)), 4))
        else:
            doc.add_paragraph(re.sub(r"[*`]", "", s).replace("|", " "))

def _add_df_table(doc, df, caption, max_rows=40):
    doc.add_heading(caption, level=3)
    show = (df.reset_index() if df.index.name else df.copy()).head(max_rows)
    t = doc.add_table(rows=1, cols=len(show.columns))
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        pass
    for j, c in enumerate(show.columns):
        t.rows[0].cells[j].text = str(c)
    for _, r in show.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(show.columns):
            cells[j].text = str(r[c])[:120]

doc = Document()
doc.add_heading("我国省域创新医药政策研究", level=0)
sub = doc.add_paragraph("——框架式质性分析与归纳式主题分析的整合（基于政策文本）")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

for title, _, _ in REPORT_SECTIONS:
    _add_markdown_block(doc, section_texts.get(title, ""))
    for fig_file, caption in figures_for(title):
        fp = os.path.join(FIG_DIR, fig_file)
        if os.path.exists(fp):
            doc.add_picture(fp, width=Inches(6.0))
            cap = doc.add_paragraph(caption); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)

doc.add_heading("附表", level=1)
_add_df_table(doc, tool_table, "表2 政策工具类型分布")
_add_df_table(doc, compare_df, "表 化学药与生物制品 vs 中医药 业务维度对照")
_add_df_table(doc, tool_compare_df, "表 化学药与生物制品 vs 中医药 政策工具对照（Rothwell）")
if not theme_df.empty:
    _add_df_table(doc, theme_df, "表 跨维度核心主题一览")
if not joint_df.empty:
    _add_df_table(doc, joint_df.reset_index(), "表 整合联合展示（主题×维度，1=贯穿）")

docx_path = os.path.join(OUTPUT_DIR, "创新医药政策研究报告.docx")
doc.save(docx_path)
print("Word 报告已生成（图表内嵌）：", docx_path)

# %% [markdown]
# ## 11. 汇总产出清单
# 运行结束后，所有结果都在 `政策分析输出/` 目录下。

# %%
print("=" * 60)
print("产出清单（目录：%s）" % OUTPUT_DIR)
print("=" * 60)
for f in sorted(glob.glob(os.path.join(OUTPUT_DIR, "**", "*"), recursive=True)):
    if os.path.isfile(f):
        print("  -", os.path.relpath(f, OUTPUT_DIR))
print("\n说明：")
print("  · 附录A_质性编码汇总表.xlsx —— 逐条编码 + 原文引证 + 政策工具类型（防杜撰证据）")
print("  · 表2_政策工具类型分布.xlsx —— 五类政策工具数量与占比")
print("  · 表2附_工具类型交叉表.xlsx —— 工具类型×省份/维度/药物类型")
print("  · 主题分析表.xlsx / 主题分析结果.json —— 跨维度归纳主题产出")
print("  · 表_整合联合展示_主题x维度.xlsx —— 两种方法整合（联合展示/三角互证）")
print("  · 表_省份x主题_覆盖.xlsx —— 省份×核心主题覆盖分布")
print("  · 省份_维度_编码矩阵.xlsx —— 比较矩阵")
print("  · 化学药生物制品_vs_中医药_对照表.xlsx —— 业务维度双路径对照")
print("  · 化药生物制品_vs_中医药_政策工具对照表.xlsx —— 政策工具(Rothwell)双路径对照")
print("  · 图/*.png —— 图1~图7（含图7整合联合展示热力图）")
print("  · 创新医药政策研究报告.md —— 研究报告（图表已内嵌）")
print("  · 创新医药政策研究报告.docx —— Word版研究报告（图表+附表内嵌，适合投稿）")
print("  · coding_cache_policy.json / tool_cache.json —— 编码与工具分类缓存（重复运行直接复用）")
